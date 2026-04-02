"""Generate a Word document comparing OpenClaw and Teammates memory systems."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Memory System Comparison', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('OpenClaw vs. Teammates')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
run.bold = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('April 2, 2026 (Revised)  |  Prepared by Scribe (Teammates PM)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()

# --- Executive Summary ---
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'This document compares the memory systems of OpenClaw and Teammates to identify '
    'improvements from Teammates that could be ported to OpenClaw. Both systems use '
    'Markdown files as the primary storage format with vector-based semantic search, '
    'but differ significantly in memory organization, lifecycle management, context '
    'budgeting, and knowledge distillation.'
)

# --- Table of Contents ---
doc.add_heading('Contents', level=1)
toc_items = [
    '1. Architecture Overview',
    '2. Side-by-Side Feature Comparison',
    '3. Memory Storage & Organization',
    '4. Memory Types & Taxonomy',
    '5. Search & Retrieval',
    '6. Context Loading & Budget Management',
    '7. Memory Lifecycle & Compaction',
    '8. Knowledge Distillation (WISDOM)',
    '9. Multi-Agent Coordination',
    '10. Recommended Improvements for OpenClaw',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# =============================================================================
# 1. Architecture Overview
# =============================================================================
doc.add_heading('1. Architecture Overview', level=1)

doc.add_heading('OpenClaw', level=2)
doc.add_paragraph(
    'OpenClaw uses a centralized memory system with a SQLite-backed vector index. '
    'Memory files live in the agent workspace (default ~/.openclaw/workspace) and are '
    'indexed into a per-agent SQLite database at ~/.openclaw/memory/{agentId}.sqlite. '
    'The system supports multiple embedding providers (OpenAI, Gemini, Voyage, Mistral, '
    'local GGUF models) and offers hybrid search combining vector similarity with BM25 '
    'full-text search. An optional QMD backend provides an alternative retrieval engine.'
)
doc.add_paragraph('Key files:')
bullets = [
    'MEMORY.md / memory.md - Main durable memory file',
    'memory/YYYY-MM-DD.md - Daily notes (append-only)',
    'Session transcripts (.jsonl) - Optional indexing of conversation history',
    'SQLite database - Chunks, embeddings, FTS5 index, embedding cache',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Teammates', level=2)
doc.add_paragraph(
    'Teammates uses a distributed, per-agent memory system where each teammate maintains '
    'its own memory folder under .teammates/<name>/memory/. Indexing uses Vectra (local '
    'vector store) with transformers.js embeddings (all-MiniLM-L6-v2, 384-dim). The '
    'system has no cloud dependencies and runs entirely locally. Memory is organized into '
    'a multi-tier hierarchy: daily logs, typed memories, weekly/monthly summaries, and '
    'a distilled WISDOM.md file.'
)
doc.add_paragraph('Key files:')
bullets = [
    'SOUL.md - Agent identity and boundaries (always loaded)',
    'GOALS.md - Active objectives (always loaded)',
    'WISDOM.md - Distilled principles (always loaded)',
    'memory/YYYY-MM-DD.md - Daily logs with YAML frontmatter',
    'memory/<type>_<topic>.md - Typed memories (user, feedback, project, reference, decision)',
    'memory/weekly/YYYY-Wnn.md - Weekly summaries (compacted from dailies)',
    'memory/monthly/YYYY-MM.md - Monthly summaries (compacted from weeklies)',
    '.index/ - Vectra vector index (gitignored)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# =============================================================================
# 2. Side-by-Side Feature Comparison
# =============================================================================
doc.add_heading('2. Side-by-Side Feature Comparison', level=1)

comparison_data = [
    ('Feature', 'OpenClaw', 'Teammates'),
    ('Storage format', 'Markdown files', 'Markdown files with YAML frontmatter'),
    ('Index backend', 'SQLite + sqlite-vec + FTS5', 'Vectra (local JSON-based vector store)'),
    ('Embedding providers', 'OpenAI, Gemini, Voyage, Mistral, local GGUF', 'Local only (all-MiniLM-L6-v2 via transformers.js)'),
    ('Cloud dependencies', 'Optional (embedding providers)', 'None (fully local)'),
    ('Memory taxonomy', 'Partial (5 categories in LanceDB backend: preference, fact, decision, entity, other; flat elsewhere)', 'Typed (user, feedback, project, reference, decision)'),
    ('Multi-tier compaction', 'No', 'Yes (daily -> weekly -> monthly)'),
    ('Knowledge distillation', 'No', 'Yes (WISDOM.md from typed memories)'),
    ('Context budgeting', 'Partial (character-based limits via QMD: maxInjectedChars 4k, maxSnippetChars 700, maxResults 6)', 'Yes (32k total, per-section allocation with overflow)'),
    ('Multi-agent memory', 'Per-agent isolated', 'Per-agent with cross-team sharing'),
    ('Session memory', 'Yes (experimental, JSONL transcripts)', 'No (daily logs serve this purpose)'),
    ('Hybrid search', 'Yes (vector + BM25, weighted)', 'Yes (vector + BM25 via Vectra, plus frontmatter catalog matching)'),
    ('MMR re-ranking', 'Yes (configurable)', 'No'),
    ('Temporal decay', 'Yes (configurable half-life)', 'Implicit (recency depth parameter)'),
    ('Query expansion', 'Yes (keyword extraction)', 'Yes (stopword filtering, query variations)'),
    ('Memory flush on compaction', 'Yes (LLM-prompted)', 'No (manual by agent)'),
    ('Batch embedding', 'Yes (OpenAI, Gemini, Voyage APIs)', 'No (local only)'),
    ('Embedding cache', 'Yes (SQLite-based, hash-keyed)', 'Implicit (Vectra dedup)'),
    ('CLI tools', 'openclaw memory status/index/search', 'teammates-recall search/index/sync/status'),
    ('Frontmatter metadata', 'No', 'Yes (name, description, type for typed memories)'),
    ('File watching', 'Yes (chokidar, debounced)', 'No (sync-on-search)'),
    ('Migration system', 'No formal system', 'Yes (versioned, in settings.json)'),
    ('Retention policies', 'Manual deletion', 'Automated (30d daily, 52w weekly, permanent monthly)'),
]

table = doc.add_table(rows=len(comparison_data), cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (feat, oc, tm) in enumerate(comparison_data):
    row = table.rows[i]
    row.cells[0].text = feat
    row.cells[1].text = oc
    row.cells[2].text = tm
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(2.0)
    row.cells[1].width = Inches(2.5)
    row.cells[2].width = Inches(2.5)

doc.add_page_break()

# =============================================================================
# 3. Memory Storage & Organization
# =============================================================================
doc.add_heading('3. Memory Storage & Organization', level=1)

doc.add_heading('OpenClaw Approach', level=2)
doc.add_paragraph(
    'OpenClaw uses a flat file structure: a single MEMORY.md for durable facts and '
    'memory/YYYY-MM-DD.md files for daily notes. There is no formal taxonomy or '
    'frontmatter convention. Memory files are plain Markdown without metadata headers. '
    'The agent decides what to write and where based on system prompt instructions.'
)

doc.add_heading('Teammates Approach', level=2)
doc.add_paragraph(
    'Teammates enforces a structured taxonomy through YAML frontmatter. Every memory file '
    'has a type field (daily, user, feedback, project, reference, decision) and typed '
    'memories include name and description fields. This metadata enables cheap catalog '
    'matching without embeddings and provides clear semantics for what each memory '
    'represents.'
)

p = doc.add_paragraph()
run = p.add_run('Improvement opportunity: ')
run.bold = True
p.add_run(
    'Adding YAML frontmatter with type/name/description to OpenClaw memory files would '
    'enable metadata-based filtering and catalog matching as a low-cost complement to '
    'vector search.'
)

# =============================================================================
# 4. Memory Types & Taxonomy
# =============================================================================
doc.add_heading('4. Memory Types & Taxonomy', level=1)

doc.add_heading('OpenClaw', level=2)
doc.add_paragraph(
    'OpenClaw has two implicit types in its core backend: durable long-term (MEMORY.md) '
    'and day-to-day context (daily notes). The newer LanceDB backend introduces 5 '
    'categories (preference, fact, decision, entity, other), providing a rudimentary '
    'taxonomy. However, these categories are backend-specific and not available in the '
    'default SQLite backend. Session transcripts are an optional third source but are '
    'experimental.'
)

doc.add_heading('Teammates', level=2)
doc.add_paragraph('Teammates defines five explicit typed memory categories:')

type_data = [
    ('Type', 'Purpose', 'Example'),
    ('user', 'User profile, preferences, knowledge level', 'User is a senior Go dev, new to React'),
    ('feedback', 'Corrections and validated approaches', 'Don\'t mock the DB in integration tests'),
    ('project', 'Ongoing work, deadlines, decisions', 'Merge freeze begins 2026-03-05'),
    ('reference', 'Pointers to external resources', 'Pipeline bugs tracked in Linear "INGEST"'),
    ('decision', 'Design decisions with reasoning', 'Chose SQLite over Postgres for local-first'),
]
table = doc.add_table(rows=len(type_data), cols=3, style='Light Grid Accent 1')
for i, (t, p_text, e) in enumerate(type_data):
    table.rows[i].cells[0].text = t
    table.rows[i].cells[1].text = p_text
    table.rows[i].cells[2].text = e
    if i == 0:
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Improvement opportunity: ')
run.bold = True
p.add_run(
    'Introducing a typed memory taxonomy in OpenClaw would help agents organize '
    'knowledge by purpose rather than dumping everything into MEMORY.md. Feedback '
    'memories are especially valuable - they prevent agents from repeating mistakes '
    'across sessions.'
)

doc.add_page_break()

# =============================================================================
# 5. Search & Retrieval
# =============================================================================
doc.add_heading('5. Search & Retrieval', level=1)

doc.add_heading('Shared Strengths', level=2)
bullets = [
    'Hybrid search (vector + BM25) - both systems combine semantic similarity with keyword matching for more robust retrieval',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('OpenClaw Strengths', level=2)
bullets = [
    'MMR re-ranking for diversity - prevents near-duplicate results',
    'Temporal decay with configurable half-life - recency-aware scoring',
    'Multiple embedding providers with fallback chains',
    'Batch embedding for large corpora (OpenAI, Gemini, Voyage)',
    'Embedding cache to avoid re-computing unchanged content',
    'File watching for near-real-time index updates',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Teammates Strengths', level=2)
bullets = [
    'Two-pass recall: automatic pre-task search + agent-driven mid-task search',
    'Frontmatter catalog matching (zero-cost metadata search before vector search)',
    'Multi-query fusion: fires primary + keyword + topic queries, deduplicates by URI',
    'Query variation generation from task prompt (no LLM needed)',
    'Fully local - no API keys, no network, no cost per query',
    'Typed memory boost - can weight typed memories higher than daily logs',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Improvement opportunities: ')
run.bold = True

improvements = [
    'Pre-task automatic recall: OpenClaw could automatically search memory before the agent starts working, injecting relevant context into the prompt without the agent needing to call memory_search.',
    'Frontmatter catalog matching: A cheap text-match pass over memory metadata before falling back to vector search. Fast, free, and catches exact-match cases that embedding similarity might miss.',
    'Multi-query fusion: Generating query variations from the user prompt and fusing results would improve recall coverage.',
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Bullet')

# =============================================================================
# 6. Context Loading & Budget Management
# =============================================================================
doc.add_heading('6. Context Loading & Budget Management', level=1)

doc.add_heading('OpenClaw', level=2)
doc.add_paragraph(
    'OpenClaw now implements character-based context limits via QMD configuration: '
    'maxInjectedChars (4,000 total cap), maxSnippetChars (700 per result), and '
    'maxResults (6 result cap). The clampResultsByInjectedChars() function in '
    'tools.citations.ts enforces these limits during memory injection. While this '
    'prevents unbounded memory consumption, it is not as granular as a per-section '
    'token budget — there is no priority-based allocation between memory types or '
    'deduplication between prompt sections.'
)

doc.add_heading('Teammates', level=2)
doc.add_paragraph(
    'Teammates implements a sophisticated context budget system (32k tokens total) with '
    'per-section allocation:'
)

budget_data = [
    ('Section', 'Budget', 'Notes'),
    ('SOUL.md (identity)', 'Outside budget', 'Always loaded, never trimmed'),
    ('GOALS.md (objectives)', 'Outside budget', 'Always loaded'),
    ('WISDOM.md (principles)', 'Outside budget', 'Always loaded, distilled to stay small'),
    ('Today\'s daily log', 'Outside budget', 'Always loaded (still being written)'),
    ('Past daily logs (days 2-7)', '12,000 tokens', 'Whole entries, newest first'),
    ('Recall results', '8,000+ tokens', 'Gets unused daily budget as overflow'),
    ('Recall overflow grace', '4,000 tokens', 'For the last entry that straddles the limit'),
    ('User profile (USER.md)', 'Within budget', 'User preferences and context'),
]
table = doc.add_table(rows=len(budget_data), cols=3, style='Light Grid Accent 1')
for i, (s, b, n) in enumerate(budget_data):
    table.rows[i].cells[0].text = s
    table.rows[i].cells[1].text = b
    table.rows[i].cells[2].text = n
    if i == 0:
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Improvement opportunity: ')
run.bold = True
p.add_run(
    'OpenClaw\'s character-based limits are a good foundation. The next step is '
    'priority-based allocation: always-loaded core files (identity, wisdom) outside the '
    'budget, with remaining budget distributed across recent daily logs and recall results. '
    'Key insight from Teammates: deduplication between "already in prompt" and "recall '
    'results" prevents wasting tokens on content the agent can already see.'
)

doc.add_page_break()

# =============================================================================
# 7. Memory Lifecycle & Compaction
# =============================================================================
doc.add_heading('7. Memory Lifecycle & Compaction', level=1)

doc.add_heading('OpenClaw', level=2)
doc.add_paragraph(
    'OpenClaw has a memory flush mechanism that prompts the agent to save durable '
    'knowledge before context compaction occurs. However, there is no automated '
    'compaction of memory files themselves. Old daily notes accumulate indefinitely '
    'unless manually deleted. The memory_search hook on /new and /reset captures '
    'session highlights, but there is no tiered retention or summarization pipeline.'
)

doc.add_heading('Teammates', level=2)
doc.add_paragraph(
    'Teammates implements a three-tier compaction pipeline with automated retention:'
)

lifecycle_data = [
    ('Tier', 'Source', 'Output', 'Retention', 'Trigger'),
    ('Daily -> Weekly', '7 daily logs', 'YYYY-Wnn.md', '30 days for dailies', 'Week completion (Sunday)'),
    ('Weekly -> Monthly', 'Weeklies > 52 weeks', 'YYYY-MM.md', '52 weeks for weeklies', 'Age threshold'),
    ('Typed -> Wisdom', 'Recurring patterns', 'WISDOM.md entries', 'Permanent', 'Manual compaction cycle'),
]
table = doc.add_table(rows=len(lifecycle_data), cols=5, style='Light Grid Accent 1')
for i, row_data in enumerate(lifecycle_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text
    if i == 0:
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Improvement opportunity: ')
run.bold = True
p.add_run(
    'OpenClaw would benefit from an automated compaction pipeline. Daily notes should '
    'be summarized into weekly digests, and old weeklies into monthlies. This prevents '
    'unbounded memory growth while preserving long-term context. The retention schedule '
    '(30d daily, 52w weekly, permanent monthly) is battle-tested in Teammates.'
)

# =============================================================================
# 8. Knowledge Distillation (WISDOM)
# =============================================================================
doc.add_heading('8. Knowledge Distillation (WISDOM)', level=1)

doc.add_paragraph(
    'This is arguably Teammates\' most distinctive memory innovation. OpenClaw has no '
    'equivalent.'
)

doc.add_heading('How WISDOM Works', level=2)
doc.add_paragraph(
    'WISDOM.md is a curated file of distilled principles extracted from typed memories. '
    'It sits outside the context budget (always loaded) and contains short, actionable '
    'heuristics - not incident reports or implementation details. The distillation process:'
)

steps = [
    'Agent reviews typed memory files looking for recurring patterns',
    'Patterns are distilled into 1-3 sentence principled rules',
    'Rules are written to WISDOM.md with a "Last compacted" timestamp',
    'Source memories that were fully absorbed are deleted',
    'Active/evolving memories are left in place',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph()
doc.add_paragraph('Example WISDOM entry:')
p = doc.add_paragraph()
p.style = 'Intense Quote'
p.add_run(
    'Prompt structure drives compliance\n'
    'Put context first, the concrete task next, and hard rules last. '
    'Restate the user request near the bottom so the model ends on '
    'the actual ask, not on background instructions.'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Improvement opportunity: ')
run.bold = True
p.add_run(
    'A WISDOM equivalent in OpenClaw would give agents a persistent, always-loaded '
    'set of learned principles. Instead of searching for past mistakes every session, '
    'the agent starts with distilled lessons. This is especially powerful for feedback '
    'memories - corrections the user gave once never need repeating.'
)

doc.add_page_break()

# =============================================================================
# 9. Multi-Agent Coordination
# =============================================================================
doc.add_heading('9. Multi-Agent Coordination', level=1)

doc.add_heading('OpenClaw', level=2)
doc.add_paragraph(
    'OpenClaw agents have isolated memory stores (per-agent SQLite databases). There is '
    'no built-in mechanism for agents to share memories, coordinate on decisions, or '
    'maintain a cross-agent knowledge base. Each agent operates in its own workspace.'
)

doc.add_heading('Teammates', level=2)
doc.add_paragraph(
    'Teammates has several multi-agent memory coordination features:'
)

bullets = [
    'Per-agent memory folders - each teammate maintains its own SOUL.md, WISDOM.md, GOALS.md, and memory/',
    'Cross-team search - any teammate can query across all teammates\' indexes',
    'CROSS-TEAM.md - shared lessons that affect multiple teammates, avoids duplication',
    'Ownership scopes - explicit table of who owns what, preventing conflicting edits',
    'Handoff context - when work is delegated, the full task context is passed in the prompt',
    'Shared docs with pointers - private docs under teammate folders, with a central index',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Improvement opportunity: ')
run.bold = True
p.add_run(
    'If OpenClaw supports multi-agent scenarios, cross-agent memory search and a shared '
    'knowledge base (equivalent to CROSS-TEAM.md) would prevent knowledge silos. Even '
    'for single-agent use, the concept of ownership scopes helps the agent understand '
    'what it should and shouldn\'t modify.'
)

# =============================================================================
# 10. Recommended Improvements for OpenClaw
# =============================================================================
doc.add_heading('10. Recommended Improvements for OpenClaw', level=1)

doc.add_paragraph(
    'Below is a prioritized list of Teammates memory innovations that would add the '
    'most value to OpenClaw, ordered by impact and implementation complexity.'
)

doc.add_heading('Priority 1 - High Impact, Moderate Effort', level=2)

improvements_p1 = [
    (
        'Typed Memory Taxonomy',
        'OpenClaw\'s LanceDB backend already defines 5 categories (preference, fact, decision, '
        'entity, other). Extend this to the core backend, add YAML frontmatter (type, name, '
        'description) to memory files, and align categories with purpose-driven types like '
        'Teammates\' user, feedback, project, reference, decision. This enables metadata '
        'filtering, catalog matching, and gives agents clear guidance on what to store where.',
        'Requires: frontmatter parser, unified type system across backends, updated agent instructions.'
    ),
    (
        'Knowledge Distillation (WISDOM.md)',
        'Add an always-loaded file of distilled principles that agents read at session start. '
        'Periodically extract recurring patterns from typed memories into short, principled rules. '
        'This gives agents persistent learned behavior without searching every session.',
        'Requires: new file convention, compaction logic, system prompt changes to always load it.'
    ),
    (
        'Context Budget Management',
        'OpenClaw already has character-based limits (maxInjectedChars, maxSnippetChars, '
        'maxResults via QMD). The next step is per-section priority allocation: always load '
        'core files (identity, wisdom) outside the budget, allocate remaining budget to recent '
        'daily logs and recall results, and deduplicate to avoid injecting the same content twice.',
        'Requires: token estimation, priority-based allocation logic, dedup between prompt sections.'
    ),
    (
        'Automated Compaction Pipeline',
        'Add daily-to-weekly and weekly-to-monthly compaction with configurable retention. '
        'Daily logs older than 30 days get summarized into weekly digests. Weeklies older '
        'than 52 weeks get compacted into monthly summaries.',
        'Requires: compaction scheduler, summary generation, retention enforcement.'
    ),
]

for title, desc, req in improvements_p1:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    run = p.add_run('Implementation: ')
    run.bold = True
    run.font.size = Pt(10)
    r2 = p.add_run(req)
    r2.font.size = Pt(10)

doc.add_heading('Priority 2 - Medium Impact, Lower Effort', level=2)

improvements_p2 = [
    (
        'Pre-Task Automatic Recall',
        'Before the agent starts working, automatically search memory using the task prompt '
        'and inject relevant results into the system prompt. This gives agents memory context '
        'without requiring them to explicitly call memory_search.',
        'Requires: query variation generation, automatic search on session/task start, prompt injection.'
    ),
    (
        'Frontmatter Catalog Matching',
        'Add a zero-cost metadata pass before vector search. Scan memory file frontmatter '
        '(name + description) for text matches against the query. This catches exact-match '
        'cases that embedding similarity might rank lower.',
        'Requires: frontmatter index (in-memory scan), merge with vector results.'
    ),
    (
        'Multi-Query Fusion',
        'Generate query variations from the user prompt (original, keyword-focused, topic-focused) '
        'and fuse results by deduplicating on file path. This improves recall coverage beyond '
        'a single query vector.',
        'Requires: query variation generator, multi-search executor, result dedup/merge.'
    ),
    (
        'Feedback Memory Loop',
        'When the user corrects the agent, automatically save a feedback-typed memory. On '
        'session start, load recent feedback memories into the prompt so corrections are '
        'never forgotten.',
        'Requires: feedback detection heuristic, typed memory creation, prompt injection.'
    ),
]

for title, desc, req in improvements_p2:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    run = p.add_run('Implementation: ')
    run.bold = True
    run.font.size = Pt(10)
    r2 = p.add_run(req)
    r2.font.size = Pt(10)

doc.add_heading('Priority 3 - Nice to Have', level=2)

improvements_p3 = [
    (
        'Cross-Agent Memory Sharing',
        'If OpenClaw supports multi-agent scenarios, enable cross-agent memory search and '
        'a shared knowledge base. Avoids knowledge silos when multiple agents work on the '
        'same project.',
    ),
    (
        'Migration System',
        'Add versioned migration tracking (like Teammates\' settings.json cliVersion) so '
        'memory format changes can be applied incrementally without data loss.',
    ),
    (
        'Identity & Goals Files',
        'Introduce SOUL.md (agent identity/boundaries) and GOALS.md (active objectives) '
        'as always-loaded context. Gives agents persistent personality and direction across '
        'sessions without consuming search budget.',
    ),
    (
        'Session Standup / Progress Summaries',
        'Generate delta-focused summaries of what changed since last session. Teammates '
        'uses daily standups that emphasize what changed rather than repeating static state.',
    ),
]

for title, desc in improvements_p3:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# =============================================================================
# Appendix: Feature Matrix
# =============================================================================
doc.add_heading('Appendix: Feature Gap Matrix', level=1)

doc.add_paragraph(
    'Quick reference showing which Teammates features are present in OpenClaw '
    'and which represent gaps.'
)

gap_data = [
    ('Teammates Feature', 'OpenClaw Status', 'Gap?'),
    ('Typed memory taxonomy (user/feedback/project/ref/decision)', 'Partial (5 categories in LanceDB backend only)', 'Partial'),
    ('YAML frontmatter on memory files', 'Not present', 'YES'),
    ('WISDOM.md (distilled principles, always loaded)', 'Not present', 'YES'),
    ('GOALS.md (active objectives, always loaded)', 'Not present', 'YES'),
    ('SOUL.md (agent identity, always loaded)', 'Partial (system prompt)', 'Partial'),
    ('Context budget management (32k, per-section)', 'Partial (character-based limits via QMD)', 'Partial'),
    ('Daily -> Weekly compaction', 'Not present', 'YES'),
    ('Weekly -> Monthly compaction', 'Not present', 'YES'),
    ('Retention policies (30d/52w/permanent)', 'Not present', 'YES'),
    ('Pre-task automatic recall', 'Not present', 'YES'),
    ('Frontmatter catalog matching', 'Not present', 'YES'),
    ('Multi-query fusion', 'Not present', 'YES'),
    ('Cross-agent memory search', 'Not present', 'YES'),
    ('Shared knowledge base (CROSS-TEAM.md)', 'Not present', 'YES'),
    ('Migration system for memory format changes', 'Not present', 'YES'),
    ('Feedback memory auto-capture', 'Not present', 'YES'),
    ('Hybrid search (vector + BM25)', 'Both present', 'No'),
    ('MMR re-ranking', 'Present', 'No'),
    ('Temporal decay', 'Present', 'No'),
    ('Multiple embedding providers', 'Present', 'No'),
    ('Embedding cache', 'Present', 'No'),
    ('File watching for index updates', 'Present', 'No'),
    ('Session transcript indexing', 'Present (experimental)', 'No'),
    ('Memory flush before compaction', 'Present', 'No'),
    ('Batch embedding APIs', 'Present', 'No'),
]

table = doc.add_table(rows=len(gap_data), cols=3, style='Light Grid Accent 1')
for i, (feat, status, gap) in enumerate(gap_data):
    table.rows[i].cells[0].text = feat
    table.rows[i].cells[1].text = status
    table.rows[i].cells[2].text = gap
    if i == 0:
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# Save
output_path = r'C:\source\teammates\.teammates\.temp\OpenClaw_vs_Teammates_Memory_Comparison.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
